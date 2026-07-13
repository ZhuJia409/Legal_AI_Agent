# -*- coding: utf-8 -*-
"""生成 Legal_AI_Agent 项目技能（skills）介绍文档到桌面。"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DESKTOP = r"C:\Users\Administrator\Desktop\Legal_AI_Agent_技能介绍.docx"

doc = Document()

# ---------- 基础样式 ----------
normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(10.5)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # 深蓝
SUB = RGBColor(0x55, 0x55, 0x55)

def set_cn_font(run, name="Microsoft YaHei"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), name)

# ---------- 封面 ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Legal_AI_Agent 项目技能介绍")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = ACCENT
set_cn_font(r)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("skills/ 目录下全部技能（Skills）的作用与使用场景速览")
r.font.size = Pt(13)
r.font.color.rgb = SUB
set_cn_font(r)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("文档由文档生成助手根据各 SKILL.md 自动整理生成")
r.font.size = Pt(9)
r.font.color.rgb = SUB
set_cn_font(r)

doc.add_paragraph()

# ---------- 总览 ----------
doc.add_heading("一、总览", level=1)
p = doc.add_paragraph()
r = p.add_run(
    "本项目 skills/ 目录共包含 20 个技能，外加一个 README 说明文件。这些技能覆盖了："
    "项目元规范与质量门禁、后端 API 与 Python 开发、前端与 UI/UX 设计、"
    "需求与设计流程方法论，以及以 LangChain / LangGraph / Deep Agents 为核心的 AI Agent 框架生态。"
)
set_cn_font(r)
p = doc.add_paragraph()
r = p.add_run(
    "它们由 AI 在开发过程中按场景自动调用（或作为编写代码前的参考规范），"
    "帮助统一代码风格、减少踩坑、保证法律智能体项目的高质量交付。"
)
set_cn_font(r)

# 分类清单
doc.add_heading("二、技能分类一览", level=1)
categories = [
    ("项目规范与质量门禁", ["verification-before-completion", "README"]),
    ("后端 API 与 Python 开发", ["api-design", "python-patterns", "python-testing"]),
    ("前端与 UI/UX 设计", ["frontend-design", "frontend-patterns", "ui-ux-pro-max", "vercel-react-best-practices"]),
    ("需求与设计流程", ["brainstorming", "documentation-lookup"]),
    ("AI Agent 框架（LangChain / LangGraph / Deep Agents）",
     ["ecosystem-primer", "langchain-fundamentals", "langchain-middleware", "langchain-rag",
      "langgraph-fundamentals", "langgraph-persistence", "langgraph-human-in-the-loop",
      "deep-agents-core", "deep-agents-orchestration", "deep-agents-memory"]),
]
for cat, items in categories:
    h = doc.add_paragraph()
    r = h.add_run("• " + cat)
    r.bold = True
    r.font.color.rgb = ACCENT
    set_cn_font(r)
    sub = doc.add_paragraph()
    r = sub.add_run("    " + "、".join(items))
    r.font.color.rgb = SUB
    set_cn_font(r)

# ---------- 各技能详述 ----------
doc.add_heading("三、各技能作用与使用场景", level=1)

def add_skill(name, purpose, cases, note=None):
    h = doc.add_heading(name, level=2)
    p = doc.add_paragraph()
    r = p.add_run("作用：")
    r.bold = True
    set_cn_font(r)
    r = p.add_run(purpose)
    set_cn_font(r)
    p = doc.add_paragraph()
    r = p.add_run("使用场景：")
    r.bold = True
    set_cn_font(r)
    for c in cases:
        li = doc.add_paragraph(style="List Bullet")
        rr = li.add_run(c)
        set_cn_font(rr)
    if note:
        p = doc.add_paragraph()
        r = p.add_run("说明：")
        r.bold = True
        r.font.color.rgb = SUB
        set_cn_font(r)
        r = p.add_run(note)
        r.font.color.rgb = SUB
        set_cn_font(r)

# 项目规范
doc.add_heading("3.1 项目规范与质量门禁", level=2)
add_skill(
    "verification-before-completion（完成前验证）",
    "一条不可协商的铁律：在任何声称“完成 / 已修复 / 通过”之前，必须先运行验证命令并确认输出。",
    ["提交代码、创建 Pull Request 或宣布任务完成前",
     "表达满意/成功之前，或委派给其它 Agent 之后",
     "所有“测试通过 / 构建成功 / Bug 已修复”类陈述"],
    "核心理念是“Evidence before claims”。它防止未经验证的虚假完成声明，维护交付可信度。",
)
add_skill(
    "README（skills 目录说明）",
    "说明 skills/ 目录的存放范围与禁止存放的内容。",
    ["了解该目录应放什么：法律检索/合同审查/类案分析等专用技能、Agent 规范、提示词模板、评估标准",
     "确认什么不该放：第三方依赖包、数据库数据/模型权重、API Key/Token 等敏感信息"],
)

# 后端
doc.add_heading("3.2 后端 API 与 Python 开发", level=2)
add_skill(
    "api-design（REST API 设计规范）",
    "REST API 设计模式全集：资源命名、状态码、分页、过滤、错误响应、版本控制与限流。",
    ["设计新的接口端点",
     "评审已有接口契约、加分页/过滤/排序",
     "实现 API 错误处理、规划版本策略",
     "构建对外或合作伙伴 API"],
    "在本项目中，案件分析与合同审查后端模块被要求优先参考此技能来统一接口路径与错误格式。",
)
add_skill(
    "python-patterns（Python 开发规范）",
    "Pythonic 写法、PEP 8、类型标注与可维护代码的工程实践。",
    ["编写新的 Python 代码",
     "审查或重构现有 Python 代码",
     "设计 Python 包/模块结构"],
    "涵盖 EAFP、上下文管理器、dataclass、并发模式、__slots__、工具链（black/ruff/mypy/pytest）等。",
)
add_skill(
    "python-testing（Python 测试模式）",
    "基于 pytest 的测试策略：TDD、夹具、Mock、参数化与覆盖率要求。",
    ["编写新代码时遵循 TDD（红-绿-重构）",
     "设计测试套件、审查测试覆盖率",
     "搭建测试基础设施"],
    "建议目标 80%+ 覆盖率，关键路径 100%；强调 Mock 外部依赖、测试边界与异常。",
)

# 前端
doc.add_heading("3.3 前端与 UI/UX 设计", level=2)
add_skill(
    "frontend-design（前端视觉设计）",
    "以“设计主理人”视角做独特、有意图的视觉设计：配色、字体、版式与签名式视觉元素。",
    ["从零构建新 UI 或重塑现有界面",
     "需要明确的视觉方向、排版与风格决策",
     "避免模板化、千篇一律的 AI 生成感设计"],
    "强调“一个大胆点即可”、克制与自我批判，反对三种常见的 AI 默认审美套路。",
)
add_skill(
    "frontend-patterns（前端开发模式）",
    "React / Next.js 的组件、状态管理、性能优化与可访问性最佳实践。",
    ["拆分/组合 React 组件（组合优于继承、复合组件、Render Props）",
     "管理状态（useState/useReducer/Zustand/Context）",
     "数据获取、表单校验、性能优化（memo/虚拟列表/代码分割）",
     "可访问性与响应式 UI"],
)
add_skill(
    "ui-ux-pro-max（UI/UX 设计智能库）",
    "超大规模设计知识库：50+ 风格、161 配色、57 字体搭配、99 条 UX 规范与 25 种图表，覆盖 10 种技术栈。",
    ["设计新页面/落地页/仪表盘/管理后台/SaaS/移动端",
     "创建或重构 UI 组件、选择配色/字体/排版/布局",
     "审查 UI 的体验、可访问性或视觉一致性"],
    "内置 search.py 脚本可按产品类型、风格、色板、图表、UX、技术栈检索，输出设计系统与实现建议。",
)
add_skill(
    "vercel-react-best-practices（Vercel React 性能最佳实践）",
    "来自 Vercel 工程的 React/Next.js 性能优化指南，按影响优先级组织。",
    ["编写、审查或重构涉及组件、Server Components、Server Actions、数据获取的 React/Next.js 代码",
     "优化包体积、消除请求瀑布、改善渲染与水合行为"],
    "完整指南在 references/full-guide.md，按需检索相关章节，优先处理高影响项。",
)

# 流程
doc.add_heading("3.4 需求与设计流程", level=2)
add_skill(
    "brainstorming（头脑风暴到设计）",
    "通过协作式对话把想法逐步变成成型的方案与规格说明，强约束“先设计、后实现”。",
    ["任何创意性工作之前：建功能、写组件、加能力、改行为",
     "实现前需要厘清用户意图、需求与设计"],
    "含硬性门禁（HARD-GATE）：未呈现设计并获用户批准，不得写代码或调用任何实现类技能。",
)
add_skill(
    "documentation-lookup（文档查询，基于 Context7）",
    "通过 Context7 MCP 获取最新、准确的库与框架官方文档，替代可能过时的训练知识。",
    ["用户问到库的配置/设置问题（如 Next.js middleware）",
     "需要依赖某库的行为写代码（如 Prisma 查询）",
     "需要 API/参考信息，或点名了具体框架/库（React、Tailwind、Supabase 等）"],
)

# AI Agent 框架
doc.add_heading("3.5 AI Agent 框架（LangChain / LangGraph / Deep Agents）", level=2)
doc.add_paragraph().add_run(
    "该族技能用于构建法律智能体的 Agent 能力。ecosystem-primer 是入口，"
    "需先据它判断该用哪一层框架，再加载对应技能。"
).italic = True

add_skill(
    "ecosystem-primer（框架选型总入口）",
    "构建任何 LangChain / LangGraph / Deep Agents 项目前最先调用的起点。",
    ["不确定该用 LangChain、LangGraph 还是 Deep Agents 时",
     "需要了解三层工具的定位、安装与环境配置",
     "确定层级后，决定下一步该加载哪个具体技能"],
    "决策顺序：需规划/文件管理/记忆/子代理 → Deep Agents；需自定义控制流 → LangGraph；单用途固定工具 → LangChain。",
)
add_skill(
    "langchain-fundamentals（LangChain 基础）",
    "使用 create_agent() 创建代理、定义工具、用中间件做人机协作与错误处理。",
    ["构建单用途、工具集固定的简单代理",
     "RAG 管道、文档问答、纯模型调用/提示词链"],
    "当前推荐用 create_agent() 而非旧式写法；覆盖工具定义、持久化(checkpointer)、结构化输出与递归上限设置。",
)
add_skill(
    "langchain-middleware（LangChain 中间件）",
    "生产级代理的中间件模式：HITL 人工审批、自定义钩子、Command 恢复与结构化输出。",
    ["在危险工具调用前暂停，等待人工批准（approve/edit/reject）",
     "拦截工具调用做日志、重试、错误处理",
     "用 Pydantic/Zod 获取类型化、校验后的响应"],
)
add_skill(
    "langchain-rag（LangChain RAG）",
    "构建检索增强生成系统的完整管道：加载、切分、嵌入、向量存储、检索与生成。",
    ["搭建任何 RAG 系统",
     "加载 PDF/网页/目录文档、文本切分、选择向量库（FAISS/Chroma/Pinecone）",
     "把检索作为 Agent 的工具对外提供"],
)
add_skill(
    "langgraph-fundamentals（LangGraph 基础）",
    "用 StateGraph 把 Agent 工作流建模为有向图：节点、边、状态、Command、Send 与流式。",
    ["需要对 Agent 编排做细粒度控制",
     "构建带分支/循环的复杂工作流",
     "需要人机协作或持久化的有状态流程"],
    "图必须先 compile() 再执行；列表字段需用 reducer 否则后写覆盖前写；注意避免无限循环。",
)
add_skill(
    "langgraph-persistence（LangGraph 持久化）",
    "持久化层：checkpointer（线程内记忆）、thread_id、时间旅行（历史回放/分支）、Store（跨线程记忆）。",
    ["需要状态持久化、记住对话",
     "需要浏览历史检查点、回放或 fork 过去状态",
     "需要跨会话/跨线程的长期记忆（用户偏好、事实）"],
    "生产环境用 PostgresSaver 而非 InMemorySaver；update_state 会穿过 reducer，需替换时用 Overwrite。",
)
add_skill(
    "langgraph-human-in-the-loop（LangGraph 人机协作）",
    "暂停图执行、把数据交给用户、用 Command(resume=...) 恢复，并处理多级错误策略。",
    ["实现人工审批/确认流程",
     "在敏感操作（如发邮件、写库）前暂停",
     "处理用户可修正的错误、并行中断"],
    "interrupt() 前须有 checkpointer 与 thread_id，且 resume 后节点会从开头重跑——非幂等副作用需放在 interrupt 之后或用 upsert。",
)
add_skill(
    "deep-agents-core（Deep Agents 核心）",
    "基于 LangChain/LangGraph 的“开箱即用”Agent 框架：规划、文件管理、子代理、记忆、HITL、技能按需加载。",
    ["需要多步规划、长上下文文件管理",
     "需要把子任务委派给专用子代理",
     "需要跨会话持久记忆或按需加载技能"],
    "通过 create_deep_agent() 配置即可，中间件（TodoList/Filesystem/SubAgent）默认内置，无需自己实现。",
)
add_skill(
    "deep-agents-orchestration（Deep Agents 编排）",
    "三大编排能力：SubAgentMiddleware（任务委派）、TodoListMiddleware（规划）、HumanInTheLoopMiddleware（审批）。",
    ["用 task 工具把工作委派给专用子代理",
     "用 write_todos 规划与跟踪多步任务",
     "对敏感操作配置人工审批"],
    "子代理是无状态的，单次调用需给完整指令；自定义子代理不会继承主代理的技能。",
)
add_skill(
    "deep-agents-memory（Deep Agents 记忆与后端）",
    "可插拔的后端策略：StateBackend（线程内临时）、StoreBackend（跨线程持久）、FilesystemBackend、CompositeBackend（混合路由）。",
    ["临时工作文件用 StateBackend",
     "本地开发需要真实磁盘访问 + 人工确认用 FilesystemBackend",
     "跨会话记忆用 StoreBackend（生产用 PostgresStore）",
     "混合存储用 CompositeBackend 按路径路由"],
    "注意：FilesystemBackend 切勿用于 Web 服务器；StoreBackend 必须配套 store 实例；生产用 PostgresStore。",
)

# ---------- 速查表 ----------
doc.add_heading("四、速查表", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, t in enumerate(["技能", "类别", "一句话定位"]):
    hdr[i].text = t
    for p in hdr[i].paragraphs:
        for rr in p.runs:
            rr.bold = True
            rr.font.size = Pt(9)
            set_cn_font(rr)

rows = [
    ("verification-before-completion", "规范/质量", "完成前必须先验证（铁律）"),
    ("README", "说明", "skills 目录存放范围说明"),
    ("api-design", "后端", "REST API 命名/状态码/分页/错误/版本/限流"),
    ("python-patterns", "后端", "Pythonic 写法与工程规范"),
    ("python-testing", "后端", "pytest / TDD / Mock / 覆盖率"),
    ("frontend-design", "前端", "独特有意图的视觉设计主导"),
    ("frontend-patterns", "前端", "React/Next.js 组件与状态模式"),
    ("ui-ux-pro-max", "前端", "大型 UI/UX 设计知识库与检索"),
    ("vercel-react-best-practices", "前端", "Vercel 性能优化优先级指南"),
    ("brainstorming", "流程", "先设计后实现（硬门禁）"),
    ("documentation-lookup", "流程", "Context7 查最新官方文档"),
    ("ecosystem-primer", "Agent", "LangChain/LangGraph/Deep Agents 选型入口"),
    ("langchain-fundamentals", "Agent", "create_agent 建简单代理"),
    ("langchain-middleware", "Agent", "HITL 审批 + 自定义中间件"),
    ("langchain-rag", "Agent", "RAG 检索增强生成管道"),
    ("langgraph-fundamentals", "Agent", "StateGraph 有状态工作流"),
    ("langgraph-persistence", "Agent", "检查点/记忆/时间旅行"),
    ("langgraph-human-in-the-loop", "Agent", "暂停-人工确认-恢复"),
    ("deep-agents-core", "Agent", "开箱即用 Agent 框架核心"),
    ("deep-agents-orchestration", "Agent", "子代理委派/规划/审批"),
    ("deep-agents-memory", "Agent", "记忆后端与持久化策略"),
]
for name, cat, desc in rows:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = cat
    cells[2].text = desc
    for c in cells:
        for p in c.paragraphs:
            for rr in p.runs:
                rr.font.size = Pt(8.5)
                set_cn_font(rr)

# 列宽
for row in table.rows:
    row.cells[0].width = Inches(2.4)
    row.cells[1].width = Inches(0.9)
    row.cells[2].width = Inches(3.5)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("— 完 —")
r.font.color.rgb = SUB
set_cn_font(r)

doc.save(DESKTOP)
print("SAVED:", DESKTOP)
