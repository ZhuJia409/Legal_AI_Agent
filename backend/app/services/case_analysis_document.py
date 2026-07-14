from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.case_analysis import (
    CaseDraftDocumentInfo,
    DocumentDraftStageResult,
    RiskLevel,
    StrategyStageResult,
)

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_MODE_LABELS = {
    "aggressive": "激进方案",
    "balanced": "稳健方案",
    "conservative": "保守方案",
}
_RISK_LABELS = {"unknown": "待评估", "low": "低", "medium": "中", "high": "高"}


@dataclass(frozen=True, slots=True)
class GeneratedCaseDocument:
    filename: str
    content_type: str
    content: bytes
    sha256: str
    generated_at: datetime

    def to_document_info(self, analysis_id: str) -> CaseDraftDocumentInfo:
        return CaseDraftDocumentInfo(
            filename=self.filename,
            size_bytes=len(self.content),
            sha256=self.sha256,
            generated_at=self.generated_at,
            download_path=f"/api/v1/case-analyses/{analysis_id}/document",
        )


class CaseAnalysisDocumentRenderer:
    """将已校验的第 7、8 阶段组装为精简草稿，不二次改写模型事实。"""

    def render(
        self,
        *,
        analysis_id: str,
        title: str | None,
        risk_level: RiskLevel,
        strategy_stage: StrategyStageResult,
        draft_stage: DocumentDraftStageResult,
    ) -> GeneratedCaseDocument:
        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)

        heading = document.add_heading("案件处理方案与文书草稿", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"案件：{title or '未命名案件'}")
        document.add_paragraph(f"整体风险：{_RISK_LABELS[risk_level]}")

        document.add_heading("一、三套方案", level=1)
        if strategy_stage.strategies:
            for strategy in strategy_stage.strategies[:3]:
                document.add_heading(_MODE_LABELS[strategy.mode], level=2)
                document.add_paragraph(strategy.summary)
                self._add_limited_list(document, "关键步骤", strategy.steps, 3)
                self._add_limited_list(document, "前提条件", strategy.prerequisites, 2)
                self._add_limited_list(document, "主要风险", strategy.risks, 2)
        else:
            document.add_paragraph("现有材料不足，三套方案待补充信息后完善。")

        document.add_heading("二、文书草稿", level=1)
        document.add_paragraph(draft_stage.draft_title or "文书类型待律师确认")
        for section in draft_stage.draft_sections[:5]:
            document.add_paragraph(section)
        missing = list(dict.fromkeys([*draft_stage.missing_information]))[:3]
        if missing:
            self._add_limited_list(document, "待补信息", missing, 3)

        document.add_heading("三、律师复核提示", level=1)
        document.add_paragraph(
            "本文书为基于已提供材料生成的草稿，不可直接提交法院或对外使用。"
            "事实、证据、法律依据、管辖、请求范围及格式必须由专业律师复核并定稿。"
        )

        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
        generated_at = datetime.now(UTC)
        safe_title = self._safe_title(title)
        return GeneratedCaseDocument(
            filename=f"{safe_title}_案件文书草稿_{analysis_id[:8]}.docx",
            content_type=DOCX_CONTENT_TYPE,
            content=content,
            sha256=hashlib.sha256(content).hexdigest(),
            generated_at=generated_at,
        )

    @staticmethod
    def _add_limited_list(
        document: Document,
        label: str,
        items: list[str],
        limit: int,
    ) -> None:
        document.add_paragraph(f"{label}：")
        for item in items[:limit]:
            document.add_paragraph(item, style="List Bullet")

    @staticmethod
    def _safe_title(title: str | None) -> str:
        normalized = re.sub(r'[\\/:*?"<>|\r\n]+', "_", (title or "未命名案件").strip())
        return normalized[:60] or "未命名案件"

